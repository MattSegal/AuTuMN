
import os
import glob
import datetime
import autumn.model
from autumn.spreadsheet import read_input_data_xls
import numpy as np
import openpyxl as xl
import tool_kit
from docx import Document
from matplotlib import pyplot, patches
import numpy
import pylab
import platform
import os


def relax_y_axis(ax):

    """
    Matplotlib's default values often place curves very close to the top
    of axes and sometimes extend down to small fractions for plots that are
    proportions. This over-rides some of these defaults, that I don't like.
    Args:
        ax: Axis with default y-limits to be revised

    Returns:
        ylims: New y-lims that look better

    """

    ylims = list(ax.get_ylim())
    if ylims[0] < ylims[1] * .75:
        ylims[0] = 0.
    else:
        ylims[0] = ylims[0] * .6
    ylims[1] = ylims[1] * 1.1

    return ylims


def find_subplot_numbers(n):

    # Find a nice number of subplots for a panel plot
    answer = find_smallest_factors_of_integer(n)
    i = 0
    while i < 10:
        if abs(answer[0] - answer[1]) > 3:
            n = n + 1
            answer = find_smallest_factors_of_integer(n)
        i = i + 1

    return answer


def find_smallest_factors_of_integer(n):

    """
    Quick method to iterate through integers to find the smallest whole number
    fractions. Written only to be called by find_subplot_numbers.

    Args:
        n: Integer to be factorised

    Returns:
        answer: The two smallest factors of the integer

    """

    answer = [1E3, 1E3]
    for i in range(1, n + 1):
        if n % i == 0 and i+(n/i) < sum(answer):
            answer = [i, n/i]
    return answer


def humanise_y_ticks(ax):

    """
    Coded by Bosco, does a few things, including rounding
    axis values to thousands, millions or billions and abbreviating
    these to single letters.

    Args:
        ax: The adapted axis

    """


    vals = list(ax.get_yticks())
    max_val = max([abs(v) for v in vals])
    if max_val < 1e3:
        return
    if max_val >= 1e3 and max_val < 1e6:
        labels = ["%.1fK" % (v/1e3) for v in vals]
    elif max_val >= 1e6 and max_val < 1e9:
        labels = ["%.1fM" % (v/1e6) for v in vals]
    elif max_val >= 1e9:
        labels = ["%.1fB" % (v/1e9) for v in vals]
    is_fraction = False
    for label in labels:
        if label[-3:-1] != ".0":
            is_fraction = True
    if not is_fraction:
        labels = [l[:-3] + l[-1] for l in labels]
    ax.set_yticklabels(labels)


def make_axes_with_room_for_legend():

    """
    Create axes for a figure with a single plot with a reasonable
    amount of space around.

    Returns:
        ax: The axes that can be plotted on

    """

    fig = pyplot.figure()
    ax = fig.add_axes([0.1, 0.1, 0.6, 0.75])
    return ax


def set_axes_props(
        ax, xlabel=None, ylabel=None, title=None, is_legend=True,
        axis_labels=None):

    frame_colour = "grey"

    # Hide top and right border of plot
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.yaxis.set_ticks_position('left')
    ax.xaxis.set_ticks_position('bottom')

    if xlabel is not None:
        ax.set_xlabel(xlabel)

    if ylabel is not None:
        ax.set_ylabel(ylabel)

    if is_legend:
        if axis_labels:
            handles, labels = ax.get_legend_handles_labels()
            leg = ax.legend(
                handles,
                axis_labels,
                bbox_to_anchor=(1.05, 1),
                loc=2,
                borderaxespad=0.,
                frameon=False,
                prop={'size': 7})
        else:
            leg = ax.legend(
                bbox_to_anchor=(1.05, 1),
                loc=2,
                borderaxespad=0.,
                frameon=False,
                prop={'size':7})
        for text in leg.get_texts():
            text.set_color(frame_colour)

    if title is not None:
        t = ax.set_title(title)
        t.set_color(frame_colour)

    for label in (ax.get_xticklabels() + ax.get_yticklabels()):
        label.set_fontname('Arial')
        label.set_fontsize(8)

    ax.tick_params(color=frame_colour, labelcolor=frame_colour)
    for spine in ax.spines.values():
        spine.set_edgecolor(frame_colour)
    ax.xaxis.label.set_color(frame_colour)
    ax.yaxis.label.set_color(frame_colour)

    humanise_y_ticks(ax)


def indices(a, func):
    return [i for (i, val) in enumerate(a) if func(val)]


def get_nice_font_size(subplot_grid):

    # Simple function to return a reasonable font size
    # as appropriate to the number of rows of subplots in the figure
    return 2. + 8. / subplot_grid[0]


def find_truncation_points(model, left_xlimit):

    # Not going to argue that the following code is the most elegant approach
    right_xlimit_index = len(model.times) - 1
    left_xlimit_index = 0
    for i in range(len(model.times)):
        if model.times[i] > left_xlimit:
            left_xlimit_index = i
            break
    return right_xlimit_index, left_xlimit_index


def find_reasonable_year_ticks(start_time, end_time):

    """
    Simple method to find some reasonably spaced x-ticks and making sure there
    aren't too many of them

    Args:
        start_time: Plotting start time
        end_time: Plotting end time

    Returns:
        xticks: List of where the x ticks should go
    """

    # If the range is divisible by 15
    if (start_time - end_time) % 15 == 0:
        xticks_any_length = numpy.arange(start_time, end_time + 15, 15)
    # Otherwise if it's divisible by 10
    elif (start_time - end_time) % 10 == 0:
        xticks_any_length = numpy.arange(start_time, end_time + 10, 10)
    # Otherwise just give up on having ticks along axis
    else:
        xticks_any_length = [start_time, end_time]

    xticks = []
    if len(xticks_any_length) > 10:
        for i in range(len(xticks_any_length)):
            if i % 2 == 0:
                xticks += [xticks_any_length[i]]
    else:
        xticks = xticks_any_length

    return xticks


def find_standard_output_styles(labels, lightening_factor=1.):

    """
    Function to find some standardised colours for the outputs we'll typically
    be reporting on - i.e. incidence, prevalence, mortality and notifications.
    Incidence is black/grey, prevalence green, mortality red and notifications blue.

    Args:
        labels: List containing strings for the outputs that colours are needed for.
        lightening_factor: Float between zero and one that specifies how much lighter to make
            the colours - with 0. being no additional lightening (black or dark green/red/blue)
            and 1. being completely lightened to reach white.

    Returns:
        colour: Colour for plotting
        indices: List of strings to be used to find the data in the data object
        yaxis_label: Unit of measurement for outcome
        title: Title for plot (so far usually a subplot)
        patch_colour: Colour half way between colour and white
    """

    colour = []
    indices = []
    yaxis_label = []
    title = []
    patch_colour = []

    if "incidence" in labels:
        colour += [(lightening_factor, lightening_factor, lightening_factor)]
        indices += ['e_inc_100k']
        yaxis_label += ['Per 100,000 per year']
        title += ["Incidence"]
    if "mortality" in labels:
        colour += [(1., lightening_factor, lightening_factor)]
        indices += ['e_mort_exc_tbhiv_100k']
        yaxis_label += ['Per 100,000 per year']
        title += ["Mortality"]
    if "prevalence" in labels:
        colour += [(lightening_factor, 0.5 + 0.5 * lightening_factor, lightening_factor)]
        indices += ['e_prev_100k']
        yaxis_label += ['Per 100,000']
        title += ["Prevalence"]
    if "notifications" in labels:
        colour += [(lightening_factor, lightening_factor, 0.5 + 0.5 * lightening_factor)]
        yaxis_label += ['']
        title += ["Notifications"]

    # Create a colour half-way between the line colour and white for patches
    for i in range(len(colour)):
        patch_colour += [[]]
        for j in range(len(colour[i])):
            patch_colour[i] += [1. - (1. - colour[i][j]) / 2.]

    return colour, indices, yaxis_label, title, patch_colour


def make_default_line_styles(n, return_all=True):

    """
    Produces a standard set of line styles that isn't adapted to
    the data being plotted.

    Args:
        n: The number of line-styles
        return_all: Whether to return all of the styles up to n or just the last one

    Returns:
        line_styles: A list of standard line-styles, or if return_all is False,
            then the single item (for methods that are iterating through plots.

    """

    # Iterate through a standard set of line styles
    for i in range(n):
        line_styles = []
        for line in ["-", ":", "-.", "--"]:
            for colour in "rbgkmcy":
                line_styles.append(line + colour)

    if return_all:
        return line_styles
    else:
        return line_styles[n-1]


def make_related_line_styles(labels, strain_or_organ):

    colours = {}
    patterns = {}
    compartment_full_names = {}
    markers = {}
    for label in labels:
        colour, pattern, compartment_full_name, marker =\
            get_line_style(label, strain_or_organ)
        colours[label] = colour
        patterns[label] = pattern
        compartment_full_names[label] = compartment_full_name
        markers[label] = marker
    return colours, patterns, compartment_full_names, markers


def get_line_style(label, strain_or_organ):

    # Unassigned groups remain black
    colour = (0, 0, 0)
    if "susceptible_vac" in label:  # susceptible_unvac remains black
        colour = (0.3, 0.3, 0.3)
    elif "susceptible_treated" in label:
        colour = (0.6, 0.6, 0.6)
    if "latent" in label:  # latent_early remains as for latent
        colour = (0, 0.4, 0.8)
    if "latent_late" in label:
        colour = (0, 0.2, 0.4)
    if "active" in label:
        colour = (0.9, 0, 0)
    elif "detect" in label:
        colour = (0, 0.5, 0)
    elif "missed" in label:
        colour = (0.5, 0, 0.5)
    if "treatment" in label:  # treatment_infect remains as for treatment
        colour = (1, 0.5, 0)
    if "treatment_noninfect" in label:
        colour = (1, 1, 0)

    pattern = get_line_pattern(label, strain_or_organ)

    category_full_name = label
    if "susceptible" in label:
        category_full_name = "Susceptible"
    if "susceptible_fully" in label:
        category_full_name = "Fully susceptible"
    elif "susceptible_vac" in label:
        category_full_name = "BCG vaccinated, susceptible"
    elif "susceptible_treated" in label:
        category_full_name = "Previously treated, susceptible"
    if "latent" in label:
        category_full_name = "Latent"
    if "latent_early" in label:
        category_full_name = "Early latent"
    elif "latent_late" in label:
        category_full_name = "Late latent"
    if "active" in label:
        category_full_name = "Active, yet to present"
    elif "detect" in label:
        category_full_name = "Detected"
    elif "missed" in label:
        category_full_name = "Missed"
    if "treatment" in label:
        category_full_name = "Under treatment"
    if "treatment_infect" in label:
        category_full_name = "Infectious under treatment"
    elif "treatment_noninfect" in label:
        category_full_name = "Non-infectious under treatment"

    if "smearpos" in label:
        category_full_name += ", \nsmear-positive"
    elif "smearneg" in label:
        category_full_name += ", \nsmear-negative"
    elif "extrapul" in label:
        category_full_name += ", \nextrapulmonary"

    if "_ds" in label:
        category_full_name += ", \nDS-TB"
    elif "_mdr" in label:
        category_full_name += ", \nMDR-TB"
    elif "_xdr" in label:
        category_full_name += ", \nXDR-TB"

    marker = ""

    return colour, pattern, category_full_name, marker


def get_line_pattern(label, strain_or_organ):

    pattern = "-"  # Default solid line
    if strain_or_organ == "organ":
        if "smearneg" in label:
            pattern = "-."
        elif "extrapul" in label:
            pattern = ":"
    elif strain_or_organ == "strain":
        if "_mdr" in label:
            pattern = '-.'
        elif "_xdr" in label:
            pattern = '.'

    return pattern


def make_plot_title(model, labels):

    try:
        if labels is model.labels:
            title = "by each individual compartment"
        elif labels is model.compartment_types \
                or labels is model.compartment_types_bystrain:
            title = "by types of compartments"
        elif labels is model.broad_compartment_types_byorgan:
            title = "by organ involvement"
        elif labels is model.broad_compartment_types \
                or labels is model.broad_compartment_types_bystrain:
            title = "by broad types of compartments"
        elif labels is model.groups["ever_infected"]:
            title = "within ever infected compartments"
        elif labels is model.groups["infected"]:
            title = "within infected compartments"
        elif labels is model.groups["active"]:
            title = "within active disease compartments"
        elif labels is model.groups["infectious"]:
            title = "within infectious compartments"
        elif labels is model.groups["identified"]:
            title = "within identified compartments"
        elif labels is model.groups["treatment"]:
            title = "within treatment compartments"
        else:
            title = "not sure"
        return title
    except:
        return ""


def create_patch_from_dictionary(dict):

    """
    Creates an array that can be used as a patch for plotting
    Args:
        dict: Dictionary with keys 'lower_limit', 'upper_limit' and 'year'
            (at least, although 'point_estimate' will also usually be there)

    Returns:
        patch_array: The patch array for plotting
    """

    patch_array = numpy.zeros(shape=(len(dict['lower_limit']) * 2, 2))
    j = 0
    for i in dict['lower_limit']:
        # Years going forwards
        patch_array[j][0] = i
        # Years going backwards
        patch_array[-(j + 1)][0] = i
        # Lower limit data going forwards
        patch_array[j][1] = dict['lower_limit'][i]
        # Upper limit data going backwards
        patch_array[-(j + 1)][1] = dict['upper_limit'][i]
        j += 1

    return patch_array


def save_png(png):

    if png is not None:
        pylab.savefig(png, dpi=300)


def open_pngs(pngs):

    operating_system = platform.system()
    if 'Windows' in operating_system:
        os.system("start " + " ".join(pngs))
    elif 'Darwin' in operating_system:
        os.system('open ' + " ".join(pngs))


def plot_populations(model, labels, values, left_xlimit, strain_or_organ, png=None):

    right_xlimit_index, left_xlimit_index = find_truncation_points(model, left_xlimit)
    colours, patterns, compartment_full_names, markers\
        = make_related_line_styles(labels, strain_or_organ)
    ax = make_axes_with_room_for_legend()
    axis_labels = []
    ax.plot(
        model.times[left_xlimit_index: right_xlimit_index],
        model.get_var_soln("population")[left_xlimit_index: right_xlimit_index],
        'k',
        label="total", linewidth=2)
    axis_labels.append("Number of persons")

    for i_plot, plot_label in enumerate(labels):
        ax.plot(
            model.times[left_xlimit_index: right_xlimit_index],
            values[plot_label][left_xlimit_index: right_xlimit_index],
            label=plot_label, linewidth=1,
            color=colours[plot_label],
            marker=markers[plot_label],
            linestyle=patterns[plot_label])
        axis_labels.append(compartment_full_names[plot_label])

    title = make_plot_title(model, labels)

    set_axes_props(ax, 'Year', 'Persons',
                   'Population, ' + title, True,
                   axis_labels)
    save_png(png)


def plot_fractions(model, values, left_xlimit, strain_or_organ, png=None, figure_number=30):

    right_xlimit_index, left_xlimit_index = find_truncation_points(model, left_xlimit)
    colours, patterns, compartment_full_names, markers\
        = make_related_line_styles(values.keys(), strain_or_organ)
    fig = pyplot.figure(figure_number)
    ax = fig.add_axes([0.1, 0.1, 0.6, 0.75])
    axis_labels = []
    for i_plot, plot_label in enumerate(values.keys()):
        ax.plot(
            model.times[left_xlimit_index: right_xlimit_index],
            values[plot_label][left_xlimit_index: right_xlimit_index],
            label=plot_label, linewidth=1,
            color=colours[plot_label],
            marker=markers[plot_label],
            linestyle=patterns[plot_label])
        axis_labels.append(compartment_full_names[plot_label])
    title = make_plot_title(model, values.keys())
    set_axes_props(ax, 'Year', 'Proportion of population',
        'Population, ' + title, True, axis_labels)
    save_png(png)


def plot_stratified_populations(model, png=None, age_or_comorbidity='age', start_time='start_time'):

    """
    Function to plot population by age group both as raw numbers and as proportions,
    both from the start of the model and using the input argument

    Args:
        model: The entire model object being interrogated
        left_xlimit: Float value representing the time to plot from for the recent plot
        png: The name of the file to be saved

    """

    if age_or_comorbidity == 'age':
        stratification = model.agegroups
    elif age_or_comorbidity == 'comorbidity':
        stratification = model.comorbidities
    else:
        raise NameError('Stratification not permitted')

    if len(stratification) < 2:
        warnings.warn('No stratification to plot')
    else:
        # Open figure
        fig = pyplot.figure()

        # Extract data
        stratified_soln, denominator = tool_kit.sum_over_compartments(model, stratification)
        stratified_fraction = tool_kit.get_fraction_soln(stratified_soln.keys(), stratified_soln, denominator)

        colours = make_default_line_styles(len(stratification), return_all=True)

        # Loop over starting from the model start and the specified starting time
        for i_time, plot_left_time in enumerate(['recent_time', start_time]):

            # Find starting times
            right_xlimit_index, left_xlimit_index \
                = find_truncation_points(model,
                                         model.inputs['model_constants'][plot_left_time])
            title_time_text = tool_kit.find_title_from_dictionary(plot_left_time)

            # Initialise some variables
            times = model.times[left_xlimit_index: right_xlimit_index]
            lower_plot_margin_count = numpy.zeros(len(times))
            upper_plot_margin_count = numpy.zeros(len(times))
            lower_plot_margin_fraction = numpy.zeros(len(times))
            upper_plot_margin_fraction = numpy.zeros(len(times))
            legd_text = []

            for i, stratum in enumerate(stratification):

                # Find numbers or fractions in that group
                stratum_count = stratified_soln[stratum][left_xlimit_index: right_xlimit_index]
                stratum_fraction = stratified_fraction[stratum][left_xlimit_index: right_xlimit_index]

                # Add group values to the upper plot range for area plot
                for j in range(len(upper_plot_margin_count)):
                    upper_plot_margin_count[j] += stratum_count[j]
                    upper_plot_margin_fraction[j] += stratum_fraction[j]

                # Plot
                ax = fig.add_subplot(2, 2, 1 + i_time)
                ax.fill_between(times, lower_plot_margin_count, upper_plot_margin_count, facecolors=colours[i][1])

                # Create proxy for legend
                ax.plot([], [], color=colours[i][1], linewidth=6)
                if age_or_comorbidity == 'age':
                    legd_text += [tool_kit.turn_strat_into_label(stratum)]
                elif age_or_comorbidity == 'comorbidity':
                    print(tool_kit.find_title_from_dictionary(stratum))
                    legd_text += [tool_kit.find_title_from_dictionary(stratum)]

                # Cosmetic changes at the end
                if i == len(stratification)-1:
                    ax.set_ylim((0., max(upper_plot_margin_count) * 1.1))
                    ax.set_xlim(int(model.times[left_xlimit_index]),
                                model.times[right_xlimit_index])
                    ax.set_title('Total numbers' + title_time_text, fontsize=8)
                    xticks = find_reasonable_year_ticks(int(model.times[left_xlimit_index]),
                                                        model.times[right_xlimit_index])
                    ax.set_xticks(xticks)
                    for axis_to_change in [ax.xaxis, ax.yaxis]:
                        for tick in axis_to_change.get_major_ticks():
                            tick.label.set_fontsize(get_nice_font_size([2]))
                    if i_time == 1:
                        ax.legend(reversed(ax.lines), reversed(legd_text), loc=2, frameon=False, fontsize=8)

                # Plot popuation proportions
                ax = fig.add_subplot(2, 2, 3 + i_time)
                ax.fill_between(times, lower_plot_margin_fraction, upper_plot_margin_fraction, facecolors=colours[i][1])

                # Cosmetic changes at the end
                if i == len(stratification)-1:
                    ax.set_ylim((0., 1.))
                    ax.set_xlim(int(model.times[left_xlimit_index]),
                                model.times[right_xlimit_index])
                    ax.set_title('Proportion of population' + title_time_text, fontsize=8)
                    xticks = find_reasonable_year_ticks(int(model.times[left_xlimit_index]),
                                                        model.times[right_xlimit_index])
                    ax.set_xticks(xticks)
                    for axis_to_change in [ax.xaxis, ax.yaxis]:
                        for tick in axis_to_change.get_major_ticks():
                            tick.label.set_fontsize(get_nice_font_size([2]))

                # Add group values to the lower plot range for next iteration
                for j in range(len(lower_plot_margin_count)):
                    lower_plot_margin_count[j] += stratum_count[j]
                    lower_plot_margin_fraction[j] += stratum_fraction[j]

        # Finish up
        fig.suptitle('Population by ' + tool_kit.find_title_from_dictionary(age_or_comorbidity),
                     fontsize=13)
        save_png(png)


def plot_outputs_against_gtb(model,
                             labels,
                             start_time,
                             end_time_str='current_time',
                             png=None,
                             country='',
                             scenario=None,
                             gtb=True,
                             figure_number=31,
                             final_run=True):

    """
    Produces the plot for the main outputs, can handle multiple scenarios (if required).
    Save as png at the end.
    Note that if running a series of scenarios, it is expected that the last scenario to
    be run will be baseline, which should have scenario set to None.

    Args:
        model: The entire model object
        labels: A list of the outputs to be plotted
        start_time: Starting time
        end_time_str: String to access end time from data
        png:
        country: Country being plotted (just need for title)
        scenario: The scenario being run, number needed for line colour

    """

    # Get standard colours for plotting GTB data against
    colour, indices, yaxis_label, title, patch_colour = \
        find_standard_output_styles(labels, lightening_factor=0.3)

    # Get the colours for the model outputs
    if scenario is None:
        # Last scenario to run should be baseline and should be run last
        # to lay a black line over the top for comparison
        output_colour = ['-k'] * len(labels)
    else:
        # Otherwise cycling through colours
        output_colour = [make_default_line_styles(scenario, False)] * len(labels)

    # Extract the plotting data of interest
    plotting_data = []
    for i in range(len(indices)):
        plotting_data += [{}]
        for j in model.inputs.original_data['tb']:
            if indices[i] in j and '_lo' in j:
                plotting_data[i]['lower_limit'] = model.inputs.original_data['tb'][j]
            elif indices[i] in j and '_hi' in j:
                plotting_data[i]['upper_limit'] = model.inputs.original_data['tb'][j]
            elif indices[i] in j:
                plotting_data[i]['point_estimate'] = model.inputs.original_data['tb'][j]

    # Truncate data to what you want to look at (rather than going back to the dawn of time)
    right_xlimit_index, left_xlimit_index = find_truncation_points(model, start_time)

    subplot_grid = find_subplot_numbers(len(labels))

    # Time to plot until
    end_time = model.inputs.model_constants[end_time_str]

    # Not sure whether we have to specify a figure number
    fig = pyplot.figure(figure_number)

    # Overall title
    fig.suptitle(country + ' model outputs', fontsize=12)

    # Truncate notification data to years of interest
    notification_data = {}
    for i in model.inputs.original_data['notifications']['c_newinc']:
        if i > start_time:
            notification_data[i] = \
                model.inputs.original_data['notifications']['c_newinc'][i]

    for i, outcome in enumerate(labels):

        ax = fig.add_subplot(subplot_grid[0], subplot_grid[1], i + 1)

        # Plot the modelled data
        ax.plot(
            model.times[left_xlimit_index: right_xlimit_index],
            model.get_var_soln(labels[i])[left_xlimit_index: right_xlimit_index],
            color=output_colour[i][1],
            linestyle=output_colour[i][0],
            linewidth=1.5)

        # This is supposed to mean if it's the last scenario, which is the baseline
        # (provided the function has been called as intended).
        if scenario is None:

            if gtb:
            # Plot the GTB data
            # Notifications are just plotted against raw reported notifications,
            # as there are no confidence intervals around these values.

                max_modelled_output = max(model.get_var_soln(labels[i])[left_xlimit_index: right_xlimit_index])

                if outcome == 'notifications':
                    ax.plot(notification_data.keys(), notification_data.values(),
                            color=colour[i], linewidth=0.5)
                    max_notifications = max(notification_data.values())
                    if max_modelled_output > max_notifications:
                        max_notifications = max_modelled_output
                    ax.set_ylim((0., max_notifications * 1.1))

                else:
                    # Central point-estimate
                    ax.plot(plotting_data[i]['point_estimate'].keys(), plotting_data[i]['point_estimate'].values(),
                            color=colour[i], linewidth=0.5)

                    # Create the patch array
                    patch_array = create_patch_from_dictionary(plotting_data[i])

                    # Create the patch image and plot it
                    patch = patches.Polygon(patch_array, color=patch_colour[i])
                    ax.add_patch(patch)

                    max_output = max(plotting_data[i]['upper_limit'].values())
                    if max_modelled_output > max_output:
                        max_output = max_modelled_output

                    # Make y-axis range extend downwards to zero
                    ax.set_ylim((0., max_output * 1.1))

            # Set x-ticks
            xticks = find_reasonable_year_ticks(start_time, end_time)
            ax.set_xticks(xticks)

            # Adjust size of labels of x-ticks
            for axis_to_change in [ax.xaxis, ax.yaxis]:
                for tick in axis_to_change.get_major_ticks():
                    tick.label.set_fontsize(get_nice_font_size(subplot_grid))

            # Add the sub-plot title with slightly larger titles than the rest of the text on the panel
            ax.set_title(title[i], fontsize=get_nice_font_size(subplot_grid) + 2.)

            # Label the y axis with the smaller text size
            ax.set_ylabel(yaxis_label[i], fontsize=get_nice_font_size(subplot_grid))

            # Get the handles, except for the last one, which plots the data
            scenario_handles = ax.lines[:-1]

            # Make some string labels for these handles
            # (this code could probably be better)
            scenario_labels = []
            for i in range(len(scenario_handles)):
                if i < len(scenario_handles) - 1:
                    scenario_labels += ['Scenario ' + str(i + 1)]
                else:
                    scenario_labels += ['Baseline']

            # Draw the legend
            ax.legend(scenario_handles,
                      scenario_labels,
                      fontsize=get_nice_font_size(subplot_grid) - 2.,
                      frameon=False)

    if final_run:

        # Save
        save_png(png)


def plot_outputs_by_age(model,
                        start_time,
                        end_time_str='current_time',
                        png=None,
                        country='',
                        scenario=None,
                        figure_number=21,
                        final_run=True):

    """
    Produces the plot for the main outputs by age, can handle multiple scenarios (if required).
    Save as png at the end.
    Note that if running a series of scenarios, it is expected that the last scenario to
    be run will be baseline, which should have scenario set to None.
    This function is a bit less flexible than plot_outputs_against_gtb, in which you can select the
    outputs you want to plot. This one is constrained to incidence and mortality (which are the only
    ones currently calculated in the model object.

    Args:
        model: The entire model object
        start_time: Starting time
        end_time_str: String to access end time from data
        png: The filename
        country: Country being plotted (just needed for title)
        scenario: The scenario being run, number needed for line colour

    """

    # Get the colours for the model outputs
    if scenario is None:
        # Last scenario to run should be baseline and should be run last
        # to lay a black line over the top for comparison
        output_colour = ['-k']
    else:
        # Otherwise cycling through colours
        output_colour = [make_default_line_styles(scenario, False)]

    # Truncate data to what you want to look at (rather than going back to the dawn of time)
    right_xlimit_index, left_xlimit_index = find_truncation_points(model, start_time)

    subplot_grid = find_subplot_numbers(len(model.agegroups) * 2 + 1)

    # Time to plot until
    end_time = model.inputs.model_constants[end_time_str]

    # Not sure whether we have to specify a figure number
    fig = pyplot.figure(figure_number)

    # Overall title
    fig.suptitle(country + ' burden by age group', fontsize=14)

    for output_no, output in enumerate(['incidence', 'mortality']):

        # Find the highest incidence value in the time period considered across all age groups
        ymax = 0.
        for agegroup in model.agegroups:
            new_ymax = max(model.get_var_soln(output + agegroup)[left_xlimit_index: right_xlimit_index])
            if new_ymax > ymax:
                ymax = new_ymax

        for i, agegroup in enumerate(model.agegroups + ['']):

            ax = fig.add_subplot(subplot_grid[0], subplot_grid[1], i + 1 + output_no * (len(model.agegroups)+1))

            # Plot the modelled data
            ax.plot(
                model.times[left_xlimit_index: right_xlimit_index],
                model.get_var_soln(output + agegroup)[left_xlimit_index: right_xlimit_index],
                color=output_colour[0][1],
                linestyle=output_colour[0][0],
                linewidth=1.5)

            # This is supposed to mean if it's the last scenario, which is the baseline
            # (provided this function has been called as intended).
            if scenario is None:

                # Set x-ticks
                xticks = find_reasonable_year_ticks(start_time, end_time)
                ax.set_xticks(xticks)

                # Adjust size of labels of x-ticks
                for axis_to_change in [ax.xaxis, ax.yaxis]:
                    for tick in axis_to_change.get_major_ticks():
                        tick.label.set_fontsize(get_nice_font_size(subplot_grid))

                # Add the sub-plot title with slightly larger titles than the rest of the text on the panel
                ax.set_title(tool_kit.capitalise_first_letter(output) + ', '
                             + tool_kit.turn_strat_into_label(agegroup), fontsize=get_nice_font_size(subplot_grid))

                # Label the y axis with the smaller text size
                if i == 0:
                    ax.set_ylabel('Per 100,000 per year', fontsize=get_nice_font_size(subplot_grid))

                # Set upper y-limit to the maximum value for any age group during the period of interest
                ax.set_ylim(bottom=0., top=ymax)

                # Get the handles, except for the last one, which plots the data
                scenario_handles = ax.lines[:-1]

                # Make some string labels for these handles
                # (this code could probably be better)
                scenario_labels = []
                for i in range(len(scenario_handles)):
                    if i < len(scenario_handles) - 1:
                        scenario_labels += ['Scenario ' + str(i + 1)]
                    else:
                        scenario_labels += ['Baseline']

                # Draw the legend
                ax.legend(scenario_handles,
                          scenario_labels,
                          fontsize=get_nice_font_size(subplot_grid) - 2.,
                          frameon=False)

    # Save
    if final_run:
        save_png(png)


def plot_flows(model, labels, png=None):

    colours, patterns, compartment_full_names\
        = write_outputs.make_related_line_styles(labels)
    ax = write_outputs.make_axes_with_room_for_legend()
    axis_labels = []
    for i_plot, plot_label in enumerate(labels):
        ax.plot(
            model.times,
            model.get_flow_soln(plot_label) / 1E3,
            label=plot_label, linewidth=1,
            color=colours[plot_label],
            linestyle=patterns[plot_label])
        axis_labels.append(compartment_full_names[plot_label])
    write_outputs.set_axes_props(ax, 'Year', 'Change per year, thousands',
                   'Aggregate flows in/out of compartment',
                   True, axis_labels)
    write_outputs.save_png(png)


def plot_scaleup_fns(model, functions, png=None,
                     start_time_str='start_time', end_time_str='',
                     parameter_type='', country=u'', figure_number=1):

    line_styles = write_outputs.make_default_line_styles(len(functions), True)
    if start_time_str == 'recent_time':
        start_time = model.inputs.model_constants[start_time_str]
    else:
        start_time = model.inputs.model_constants[start_time_str]
    end_time = model.inputs.model_constants[end_time_str]
    x_vals = numpy.linspace(start_time, end_time, 1E3)

    pyplot.figure(figure_number)

    ax = write_outputs.make_axes_with_room_for_legend()
    for figure_number, function in enumerate(functions):
        ax.plot(x_vals,
                map(model.scaleup_fns[function],
                    x_vals), line_styles[figure_number],
                label=function)

    plural = ''
    if len(functions) > 1:
        plural += 's'
    title = str(country) + ' ' + \
            tool_kit.find_title_from_dictionary(parameter_type) + \
            ' parameter' + plural + tool_kit.find_title_from_dictionary(start_time_str)
    write_outputs.set_axes_props(ax, 'Year', 'Parameter value',
                   title, True, functions)

    ylims = write_outputs.relax_y_axis(ax)
    ax.set_ylim(bottom=ylims[0], top=ylims[1])

    write_outputs.save_png(png)


def plot_comparative_age_parameters(data_strat_list,
                                    data_value_list,
                                    model_value_list,
                                    model_strat_list,
                                    parameter_name):

    # Get good tick labels from the stratum lists
    data_strat_labels = []
    for i in range(len(data_strat_list)):
        data_strat_labels += [tool_kit.turn_strat_into_label(data_strat_list[i])]
    model_strat_labels = []
    for i in range(len(model_strat_list)):
        model_strat_labels += [tool_kit.turn_strat_into_label(model_strat_list[i])]

    # Find a reasonable upper limit for the y-axis
    ymax = max(data_value_list + model_value_list) * 1.2

    # Plot original data bar charts
    subplot_grid = (1, 2)
    fig = pyplot.figure()
    ax = fig.add_axes([0.1, 0.2, 0.35, 0.6])
    x_positions = range(len(data_strat_list))
    width = .6
    ax.bar(x_positions, data_value_list, width)
    ax.set_ylabel('Parameter value',
                  fontsize=write_outputs.get_nice_font_size(subplot_grid))
    ax.set_title('Input data', fontsize=12)
    ax.set_xticklabels(data_strat_labels, rotation=45)
    ax.set_xticks(x_positions)
    ax.set_ylim(0., ymax)
    ax.set_xlim(-1. + width, x_positions[-1] + 1)

    # Plot adjusted parameters bar charts
    ax = fig.add_axes([0.55, 0.2, 0.35, 0.6])
    x_positions = range(len(model_strat_list))
    ax.bar(x_positions, model_value_list, width)
    ax.set_title('Model implementation', fontsize=12)
    ax.set_xticklabels(model_strat_labels, rotation=45)
    ax.set_xticks(x_positions)
    ax.set_ylim(0., ymax)
    ax.set_xlim(-1. + width, x_positions[-1] + 1)

    # Overall title
    fig.suptitle(tool_kit.capitalise_first_letter(tool_kit.replace_underscore_with_space(parameter_name))
                 + ' adjustment',
                 fontsize=15)

    # Find directory and save
    out_dir = 'fullmodel_graphs'
    if not os.path.isdir(out_dir):
        os.makedirs(out_dir)
    base = os.path.join(out_dir, parameter_name)
    write_outputs.save_png(base + '_param_adjustment.png')



class Project:

    def __init__(self, country, inputs):

        """
        Initialises an object of class Project, that will contain all the information (data + outputs) for writing a
        report for a country
        Args:
            models: dictionary such as: models = {'baseline': model, 'scenario_1': model_1,  ...}
        """

        self.country = country.lower()
        self.name = 'test_' + self.country
        self.scenarios = []
        self.models = {}
        self.full_output_dict = {}
        self.integer_output_dict = {}
        self.inputs = inputs
        self.out_dir_project = os.path.join('projects', self.name)
        if not os.path.isdir(self.out_dir_project):
            os.makedirs(self.out_dir_project)

    #################################
    # General methods for use below #
    #################################

    def find_years_to_write(self, scenario, output, minimum=0, maximum=3000, step=1):

        requested_years = range(minimum, maximum, step)
        years = []
        for y in self.integer_output_dict[scenario][output].keys():
            if y in requested_years:
                years += [y]
        return years

    #########################################
    # Methods to collect data for later use #
    #########################################

    def create_output_dicts(self, outputs=['incidence', 'mortality', 'prevalence', 'notifications']):

        """
        Works through all the methods to this object that are required to populate the output dictionaries.
        First the "full" ones with all time point included, then the abbreviated ones.

        Args:
            outputs: The outputs to be populated to the dictionaries
        """

        self.create_full_output_dict(outputs)
        self.add_full_economics_dict()
        self.extract_integer_dict()

    def create_full_output_dict(self, outputs):

        """
        Creates a dictionary for each requested output at every time point in that model's times attribute
        """

        for scenario in self.scenarios:
            self.full_output_dict[scenario] = {}
            for label in outputs:
                times = self.models[scenario].times
                solution = self.models[scenario].get_var_soln(label)
                self.full_output_dict[scenario][label] = dict(zip(times, solution))

    def add_full_economics_dict(self):

        """
        Creates an economics dictionary structure that mirrors that of the epi dictionaries and adds
        this to the main outputs (epi) dictionary
        """

        for model in self.models:
            economics_dict = {}
            for intervention in self.models[model].costs:
                if intervention != 'cost_times':
                    economics_dict['cost_' + intervention] = {}
                    for t in range(len(self.models[model].costs['cost_times'])):
                        economics_dict['cost_' + intervention][self.models[model].costs['cost_times'][t]] \
                            = self.models[model].costs[intervention]['raw_cost'][t]
            self.full_output_dict[model].update(economics_dict)

    def extract_integer_dict(self):

        """
        Extracts a dictionary from full_output_dict with only integer years, using the first time value greater than
        the integer year in question.
        """

        for model in self.models:
            self.integer_output_dict[model] = {}
            for output in self.full_output_dict[model]:
                self.integer_output_dict[model][output] = {}
                times = self.full_output_dict[model][output].keys()
                times.sort()
                start = np.floor(times[0])
                finish = np.floor(times[-1])
                float_years = np.linspace(start, finish, finish - start + 1.)
                for year in float_years:
                    key = [t for t in times if t >= year][0]
                    self.integer_output_dict[model][output][int(key)] \
                        = self.full_output_dict[model][output][key]

    #################################################
    # Methods for outputting to Office applications #
    #################################################

    def write_spreadsheets(self):

        """
        Determine whether to write to spreadsheets by scenario or by output
        """

        if self.inputs.model_constants['output_spreadsheets']:
            if self.inputs.model_constants['output_by_scenario']:
                print('Writing scenario spreadsheets')
                self.write_xls_by_scenario()
            else:
                print('Writing output indicator spreadsheets')
                self.write_xls_by_output()

    def write_xls_by_output(self):

        # Find directory to write to
        out_dir_project = self.find_or_make_directory()

        # Write a new file for each output
        outputs = self.integer_output_dict['baseline'].keys()
        for output in outputs:

            # Make filename
            path = os.path.join(self.out_dir_project, output)
            path += ".xlsx"

            # Get active sheet
            wb = xl.Workbook()
            sheet = wb.active
            sheet.title = output

            # Write a new file for each epidemiological indicator
            for scenario in self.integer_output_dict.keys():
                self.write_xls_column_or_row(sheet, scenario, output)

            # Save workbook
            wb.save(path)

    def write_xls_column_or_row(self, sheet, scenario, output):

        # Find years to write
        years = self.find_years_to_write(scenario,
                                         output,
                                         int(self.inputs.model_constants['report_start_time']),
                                         int(self.inputs.model_constants['report_end_time']),
                                         int(self.inputs.model_constants['report_step_time']))

        # Write data
        if self.inputs.model_constants['output_horizontally']:
            self.write_horizontally_by_scenario(sheet, output, years)
        else:
            self.write_vertically_by_scenario(sheet, output, years)

    def write_xls_by_scenario(self):

        # Write a new file for each scenario
        scenarios = self.integer_output_dict.keys()
        for scenario in scenarios:

            # Make filename
            path = os.path.join(self.out_dir_project, scenario)
            path += '.xlsx'

            # Get active sheet
            wb = xl.Workbook()
            sheet = wb.active
            sheet.title = scenario

            for output in self.integer_output_dict['baseline'].keys():
                self.write_xls_column_or_row(sheet, scenario, output)

            # Save workbook
            wb.save(path)

    def write_horizontally_by_scenario(self, sheet, output, years):

        sheet.cell(row=1, column=1).value = 'Year'

        col = 1
        for y in years:
            col += 1
            sheet.cell(row=1, column=col).value = y

        r = 1
        for scenario in self.scenarios:
            r += 1
            sheet.cell(row=r, column=1).value = \
                tool_kit.replace_underscore_with_space(
                    tool_kit.capitalise_first_letter(scenario))
            col = 1
            for y in years:
                col += 1
                if y in self.integer_output_dict[scenario][output]:
                    sheet.cell(row=r, column=col).value \
                        = self.integer_output_dict[scenario][output][y]

    def write_horizontally_by_output(self, sheet, scenario, years):

        sheet.cell(row=1, column=1).value = 'Year'

        col = 1
        for y in years:
            col += 1
            sheet.cell(row=1, column=col).value = y

        r = 1
        for output in self.integer_output_dict['baseline'].keys():
            r += 1
            sheet.cell(row=r, column=1).value = \
                tool_kit.replace_underscore_with_space(
                    tool_kit.capitalise_first_letter(output))
            col = 1
            for y in years:
                col += 1
                if y in self.integer_output_dict[scenario][output]:
                    sheet.cell(row=r, column=col).value \
                        = self.integer_output_dict[scenario][output][y]

    def write_vertically_by_scenario(self, sheet, output, years):

        sheet.cell(row=1, column=1).value = 'Year'

        row = 1
        for y in years:
            row += 1
            sheet.cell(row=row, column=1).value = y

        col = 1
        for scenario in self.scenarios:
            col += 1
            sheet.cell(row=1, column=col).value = \
                tool_kit.replace_underscore_with_space(
                    tool_kit.capitalise_first_letter(scenario))
            row = 1
            for y in years:
                row += 1
                if y in self.integer_output_dict[scenario][output]:
                    sheet.cell(row=row, column=col).value = self.integer_output_dict[scenario][output][y]

    def write_vertically_by_output(self, sheet, scenario, years):

        sheet.cell(row=1, column=1).value = 'Year'

        row = 1
        for y in years:
            row += 1
            sheet.cell(row=row, column=1).value = y

        col = 1
        for output in self.integer_output_dict['baseline'].keys():
            col += 1
            sheet.cell(row=1, column=col).value = \
                tool_kit.replace_underscore_with_space(
                    tool_kit.capitalise_first_letter(output))
            row = 1
            for y in years:
                row += 1
                if y in self.integer_output_dict[scenario][output]:
                    sheet.cell(row=row, column=col).value \
                        = self.integer_output_dict[scenario][output][y]

    def write_documents(self):

        """
        Determine whether to write to documents by scenario or by output
        """

        if self.inputs.model_constants['output_documents']:
            if self.inputs.model_constants['output_by_scenario']:
                print('Writing scenario documents')
                self.write_docs_by_scenario()
            else:
                print('Writing output indicator documents')
                self.write_docs_by_output()

    def write_docs_by_output(self):

        # Write a new file for each output
        outputs = self.integer_output_dict['baseline'].keys()

        for output in outputs:

            # Initialise document
            path = os.path.join(self.out_dir_project, output)
            path += ".docx"
            document = Document()
            table = document.add_table(rows=1, cols=len(self.scenarios) + 1)

            # Write headers
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Year'
            for scenario_no, scenario in enumerate(self.scenarios):
                header_cells[scenario_no + 1].text \
                    = tool_kit.capitalise_first_letter(tool_kit.replace_underscore_with_space(scenario))

            # Find years to write
            years = self.find_years_to_write('baseline',
                                             output,
                                             int(self.inputs.model_constants['report_start_time']),
                                             int(self.inputs.model_constants['report_end_time']),
                                             int(self.inputs.model_constants['report_step_time']))

            for year in years:

                # Add row to table
                row_cells = table.add_row().cells
                row_cells[0].text = str(year)

                for sc, scenario in enumerate(self.scenarios):
                    if year in self.integer_output_dict[scenario][output]:
                        row_cells[sc + 1].text = '%.2f' % self.integer_output_dict[scenario][output][year]

            # Save document
            document.save(path)

    def write_docs_by_scenario(self):

        # Write a new file for each output
        outputs = self.integer_output_dict['baseline'].keys()

        for scenario in self.scenarios:

            # Initialise document
            path = os.path.join(self.out_dir_project, scenario)
            path += ".docx"
            document = Document()
            table = document.add_table(rows=1, cols=len(outputs) + 1)

            # Write headers
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Year'
            for output_no, output in enumerate(outputs):
                header_cells[output_no + 1].text \
                    = tool_kit.capitalise_first_letter(tool_kit.replace_underscore_with_space(output))

            # Find years to write
            years = self.find_years_to_write(scenario,
                                             output,
                                             int(self.inputs.model_constants['report_start_time']),
                                             int(self.inputs.model_constants['report_end_time']),
                                             int(self.inputs.model_constants['report_step_time']))

            for year in years:

                # Add row to table
                row_cells = table.add_row().cells
                row_cells[0].text = str(year)

                for out, output in enumerate(outputs):
                    if year in self.integer_output_dict[scenario][output]:
                        row_cells[out + 1].text = '%.2f' % self.integer_output_dict[scenario][output][year]

            # Save document
            document.save(path)

    def run_plotting(self):

        # Plot scale-up functions - currently only doing this for the baseline model run
        if self.inputs.model_constants['output_scaleups']:
            self.plot_classified_scaleups(self.models['baseline'])

    def plot_classified_scaleups(self, model):

        # Classify scale-up functions
        classifications = ['demo_', 'econ_', 'epi_', 'program_prop_', 'program_timeperiod']
        classified_scaleups = {}
        for classification in classifications:
            classified_scaleups[classification] = []
            for fn in model.scaleup_fns:
                if classification in fn:
                    classified_scaleups[classification] += [fn]

        base = os.path.join(self.out_dir_project, self.country + '_baseline_')

        # Time periods to perform the plots over
        times_to_plot = ['start_', 'recent_']

        # Plot them from the start of the model and from "recent_time"
        for c, classification in enumerate(classified_scaleups):
            if len(classified_scaleups[classification]) > 0:
                for j, start_time in enumerate(times_to_plot):
                    self.plot_all_scaleup_fns_against_data(model,
                                                           classified_scaleups[classification],
                                                           base + classification + '_datascaleups_from' + start_time[:-1] + '.png',
                                                           start_time + 'time',
                                                           'current_time',
                                                           classification,
                                                           figure_number=c + j * len(classified_scaleups) + 2)
                    if classification == 'program_prop':
                        plot_scaleup_fns(model,
                                                         classified_scaleups[classification],
                                                         base + classification + 'scaleups_from' + start_time[:-1] + '.png',
                                                         start_time + 'time',
                                                         'current_time',
                                                         classification,
                                                         figure_number=c + j * len(classified_scaleups) + 2 + len(classified_scaleups) * len(times_to_plot))

    def plot_all_scaleup_fns_against_data(self, model, functions, png=None,
                                          start_time_str='start_time',
                                          end_time_str='',
                                          parameter_type='',
                                          scenario=None,
                                          figure_number=2):

        # Get the colours for the model outputs
        if scenario is None:
            # Last scenario to run should be baseline and should be run last
            # to lay a black line over the top for comparison
            output_colour = ['k'] * len(functions)
        else:
            # Otherwise cycling through colours
            output_colour = [make_default_line_styles(scenario, False)[1]] * len(functions)

        # Determine how many subplots to have
        subplot_grid = find_subplot_numbers(len(functions))

        # Set x-values
        if start_time_str == 'recent_time':
            start_time = model.inputs.model_constants[start_time_str]
        else:
            start_time = model.inputs.model_constants[start_time_str]
        end_time = model.inputs.model_constants[end_time_str]
        x_vals = numpy.linspace(start_time, end_time, 1E3)

        # Initialise figure
        fig = pyplot.figure(figure_number)

        # Upper title for whole figure
        plural = ''
        if len(functions) > 1:
            plural += 's'
        title = model.inputs.model_constants['country'] + ' ' + \
                tool_kit.find_title_from_dictionary(parameter_type) + \
                ' parameter' + plural + tool_kit.find_title_from_dictionary(start_time_str)
        fig.suptitle(title)

        # Iterate through functions
        for figure_number, function in enumerate(functions):

            # Initialise subplot areas
            ax = fig.add_subplot(subplot_grid[0], subplot_grid[1], figure_number + 1)

            # Line plot scaling parameters
            ax.plot(x_vals,
                    map(model.scaleup_fns[function],
                        x_vals),
                    # line_styles[i],
                    # label=function,
                    color=output_colour[figure_number])

            if scenario is None:
                data_to_plot = {}
                for j in model.inputs.scaleup_data[scenario][function]:
                    if j > start_time:
                        data_to_plot[j] = model.inputs.scaleup_data[scenario][function][j]

                # Scatter plot data from which they are derived
                ax.scatter(data_to_plot.keys(),
                           data_to_plot.values(),
                           color=output_colour[figure_number],
                           s=6)

                # Adjust tick font size
                ax.set_xticks([start_time, end_time])
                for axis_to_change in [ax.xaxis, ax.yaxis]:
                    for tick in axis_to_change.get_major_ticks():
                        tick.label.set_fontsize(get_nice_font_size(subplot_grid))

                # Truncate parameter names depending on whether it is a
                # treatment success/death proportion
                title = tool_kit.find_title_from_dictionary(function)
                ax.set_title(title, fontsize=get_nice_font_size(subplot_grid))

                ylims = relax_y_axis(ax)
                ax.set_ylim(bottom=ylims[0], top=ylims[1])

        fig.suptitle('Scale-up functions')

        save_png(png)
